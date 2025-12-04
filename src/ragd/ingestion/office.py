"""Office document extractors (F-100).

Provides extraction for DOCX, XLSX, and EPUB formats.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from ragd.ingestion.extractor import ExtractionResult


class DOCXExtractor:
    """Extract text from Word documents.

    Requires python-docx package.
    """

    def extract(self, path: Path) -> ExtractionResult:
        """Extract text from DOCX file.

        Args:
            path: Path to DOCX file

        Returns:
            ExtractionResult with text
        """
        try:
            from docx import Document
        except ImportError:
            return ExtractionResult(
                text="",
                extraction_method="python-docx",
                success=False,
                error="python-docx not installed. Install with: pip install python-docx",
            )

        try:
            doc = Document(path)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            text = "\n\n".join(paragraphs)

            return ExtractionResult(
                text=text,
                metadata={
                    "source": str(path),
                    "format": "docx",
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                },
                extraction_method="python-docx",
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                extraction_method="python-docx",
                success=False,
                error=str(e),
            )


class XLSXExtractor:
    """Extract text from Excel spreadsheets.

    Requires openpyxl package.
    """

    def extract(self, path: Path) -> ExtractionResult:
        """Extract text from XLSX file.

        Args:
            path: Path to XLSX file

        Returns:
            ExtractionResult with text
        """
        try:
            from openpyxl import load_workbook
        except ImportError:
            return ExtractionResult(
                text="",
                extraction_method="openpyxl",
                success=False,
                error="openpyxl not installed. Install with: pip install openpyxl",
            )

        try:
            wb = load_workbook(path, read_only=True, data_only=True)
            sheets_text = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows = []

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_values):
                        rows.append(" | ".join(row_values))

                if rows:
                    sheet_text = f"## {sheet_name}\n\n" + "\n".join(rows)
                    sheets_text.append(sheet_text)

            wb.close()

            text = "\n\n".join(sheets_text)

            return ExtractionResult(
                text=text,
                metadata={
                    "source": str(path),
                    "format": "xlsx",
                    "sheets": len(wb.sheetnames),
                },
                extraction_method="openpyxl",
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                extraction_method="openpyxl",
                success=False,
                error=str(e),
            )


class EPUBExtractor:
    """Extract text from EPUB ebooks.

    Requires ebooklib package.
    """

    def extract(self, path: Path) -> ExtractionResult:
        """Extract text from EPUB file.

        Args:
            path: Path to EPUB file

        Returns:
            ExtractionResult with text
        """
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
        except ImportError:
            return ExtractionResult(
                text="",
                extraction_method="ebooklib",
                success=False,
                error="ebooklib not installed. Install with: pip install ebooklib",
            )

        try:
            book = epub.read_epub(str(path))
            chapters = []
            title = book.get_metadata("DC", "title")
            authors = book.get_metadata("DC", "creator")

            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), "html.parser")
                    text = soup.get_text(separator="\n")
                    if text.strip():
                        chapters.append(text.strip())

            text = "\n\n".join(chapters)

            metadata = {
                "source": str(path),
                "format": "epub",
                "chapters": len(chapters),
            }

            if title:
                metadata["title"] = title[0][0]
            if authors:
                metadata["creator"] = "; ".join(a[0] for a in authors)

            return ExtractionResult(
                text=text,
                metadata=metadata,
                extraction_method="ebooklib",
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                extraction_method="ebooklib",
                success=False,
                error=str(e),
            )


def get_office_extractor(path: Path) -> DOCXExtractor | XLSXExtractor | EPUBExtractor | None:
    """Get appropriate extractor for office file types.

    Args:
        path: Path to file

    Returns:
        Extractor instance or None if not supported
    """
    suffix = path.suffix.lower()

    extractors = {
        ".docx": DOCXExtractor,
        ".xlsx": XLSXExtractor,
        ".epub": EPUBExtractor,
    }

    extractor_class = extractors.get(suffix)
    return extractor_class() if extractor_class else None
